from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import utils
from docx import Document
from loguru import logger

MARGIN_X = 50
MARGIN_Y = inch
LINE_HEIGHT = 20
DEFAULT_FONT_SIZE = 12
FONT_NAME = "Helvetica"


def convert(input_path: str, output_path: str) -> str:
    """Convert DOCX file to PDF."""
    logger.info("Converting {} to PDF", input_path)

    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"File not found: {input_path}")

    doc = Document(str(input_path))

    if not doc.paragraphs:
        raise ValueError("Document has no content")

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 1 * inch
    available_width = width - 2 * MARGIN_X

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            y -= LINE_HEIGHT
            if y < MARGIN_Y:
                c.showPage()
                y = height - 1 * inch
            continue

        font_size = DEFAULT_FONT_SIZE
        for run in paragraph.runs:
            if run.font.size:
                font_size = run.font.size.pt
                break

        c.setFont(FONT_NAME, font_size)

        wrapped = utils.simpleSplit(text, FONT_NAME, font_size, available_width)
        for line in wrapped:
            if y < MARGIN_Y:
                c.showPage()
                y = height - 1 * inch
                c.setFont(FONT_NAME, font_size)
            c.drawString(MARGIN_X, y, line)
            y -= LINE_HEIGHT

    c.save()
    logger.info("Saved PDF to: {}", output_path)
    return output_path
